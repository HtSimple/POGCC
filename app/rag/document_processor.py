from __future__ import annotations

import re
import uuid
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from .schemas import RawDocument, TextChunk


@dataclass
class ProcessorConfig:
    max_chars: int = 1200
    min_chunk_chars: int = 80
    merge_short_paragraphs: bool = True
    preserve_page_info: bool = True
    pdf_page_level_chunking: bool = True


class DocumentProcessor:
    def __init__(self, config: Optional[ProcessorConfig] = None) -> None:
        self.config = config or ProcessorConfig()

    def process_file(self, file_path: str | Path) -> tuple[RawDocument, List[TextChunk]]:
        document = self.parse_file(file_path)
        chunks = self.chunk_document(document)
        return document, chunks

    def parse_file(self, file_path: str | Path) -> RawDocument:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(file_path)
        if suffix in {".docx", ".doc"}:
            return self._parse_docx(file_path)
        if suffix in {".txt", ".md"}:
            return self._parse_txt(file_path)

        raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, file_path: Path) -> RawDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        page_texts = []

        for i, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            page_text = self.clean_text(page_text)
            page_texts.append(
                {
                    "page_num": i,
                    "text": page_text,
                }
            )

        full_text = "\n\n".join(p["text"] for p in page_texts if p["text"].strip())

        return RawDocument(
            doc_id=self._make_doc_id(file_path),
            file_name=file_path.name,
            file_type="pdf",
            text=full_text,
            metadata={
                "num_pages": len(page_texts),
                "title": None,
                "source_path": str(file_path),
                "page_texts": page_texts,
            },
        )

    def _parse_docx(self, file_path: Path) -> RawDocument:
        from docx import Document

        doc = Document(str(file_path))
        parts: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        text = self.clean_text("\n\n".join(parts))

        return RawDocument(
            doc_id=self._make_doc_id(file_path),
            file_name=file_path.name,
            file_type="docx",
            text=text,
            metadata={
                "title": None,
                "source_path": str(file_path),
            },
        )

    def _parse_txt(self, file_path: Path) -> RawDocument:
        encodings = ["utf-8", "utf-8-sig", "gb18030", "latin-1"]
        text = None

        for enc in encodings:
            try:
                text = file_path.read_text(encoding=enc)
                break
            except Exception:
                continue

        if text is None:
            raise ValueError(f"Failed to read text file: {file_path}")

        text = self.clean_text(text)

        return RawDocument(
            doc_id=self._make_doc_id(file_path),
            file_name=file_path.name,
            file_type=file_path.suffix.lower().lstrip("."),
            text=text,
            metadata={
                "title": None,
                "source_path": str(file_path),
            },
        )

    def clean_text(self, text: str) -> str:
        if not text:
            return ""

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = text.replace("\u00a0", " ").replace("\ufeff", "")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        return text.strip()

    def chunk_document(self, document: RawDocument) -> List[TextChunk]:
        if document.file_type == "pdf" and self.config.pdf_page_level_chunking:
            return self._chunk_pdf_by_page(document)
        return self._chunk_normal_document(document)

    def _chunk_pdf_by_page(self, document: RawDocument) -> List[TextChunk]:
        chunks: List[TextChunk] = []
        chunk_index = 0

        page_texts = document.metadata.get("page_texts", [])
        page_texts = self._remove_repeated_page_lines(page_texts)

        for page in page_texts:
            page_num = page["page_num"]
            page_text = self.clean_text(page["text"])
            if not page_text:
                continue

            paragraphs = self._page_to_paragraphs(page_text)

            # 一页内容不多：整页一个 chunk
            if len(page_text) <= self.config.max_chars:
                chunks.append(
                    TextChunk(
                        chunk_id=f"{document.doc_id}_chunk_{chunk_index:04d}",
                        doc_id=document.doc_id,
                        text=page_text,
                        chunk_index=chunk_index,
                        source_file=document.file_name,
                        section_title=None,
                        start_page=page_num,
                        end_page=page_num,
                        pages=[page_num],
                        metadata={
                            "file_type": document.file_type,
                            "num_pages": document.metadata.get("num_pages"),
                            "title": document.metadata.get("title"),
                            "source_path": document.metadata.get("source_path"),
                        },
                    )
                )
                chunk_index += 1
                continue

            # 一页内容过多：只在当前页内部继续切，不跨页
            buffer: List[str] = []
            buffer_len = 0

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                if buffer and buffer_len + len(para) > self.config.max_chars:
                    chunk_text = "\n".join(buffer).strip()
                    chunks.append(
                        TextChunk(
                            chunk_id=f"{document.doc_id}_chunk_{chunk_index:04d}",
                            doc_id=document.doc_id,
                            text=chunk_text,
                            chunk_index=chunk_index,
                            source_file=document.file_name,
                            section_title=None,
                            start_page=page_num,
                            end_page=page_num,
                            pages=[page_num],
                            metadata={
                                "file_type": document.file_type,
                                "num_pages": document.metadata.get("num_pages"),
                                "title": document.metadata.get("title"),
                                "source_path": document.metadata.get("source_path"),
                            },
                        )
                    )
                    chunk_index += 1
                    buffer = []
                    buffer_len = 0

                buffer.append(para)
                buffer_len += len(para)

            if buffer:
                chunk_text = "\n".join(buffer).strip()
                chunks.append(
                    TextChunk(
                        chunk_id=f"{document.doc_id}_chunk_{chunk_index:04d}",
                        doc_id=document.doc_id,
                        text=chunk_text,
                        chunk_index=chunk_index,
                        source_file=document.file_name,
                        section_title=None,
                        start_page=page_num,
                        end_page=page_num,
                        pages=[page_num],
                        metadata={
                            "file_type": document.file_type,
                            "num_pages": document.metadata.get("num_pages"),
                            "title": document.metadata.get("title"),
                            "source_path": document.metadata.get("source_path"),
                        },
                    )
                )
                chunk_index += 1

        return chunks

    def _chunk_normal_document(self, document: RawDocument) -> List[TextChunk]:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", document.text) if p.strip()]
        chunks: List[TextChunk] = []
        chunk_index = 0
        buffer: List[str] = []
        buffer_len = 0

        for para in paragraphs:
            if self.config.merge_short_paragraphs and len(para) < self.config.min_chunk_chars and buffer:
                buffer.append(para)
                buffer_len += len(para)
                continue

            if buffer and buffer_len + len(para) > self.config.max_chars:
                chunk_text = "\n".join(buffer).strip()
                chunks.append(
                    TextChunk(
                        chunk_id=f"{document.doc_id}_chunk_{chunk_index:04d}",
                        doc_id=document.doc_id,
                        text=chunk_text,
                        chunk_index=chunk_index,
                        source_file=document.file_name,
                        section_title=None,
                        start_page=None,
                        end_page=None,
                        pages=[],
                        metadata={
                            "file_type": document.file_type,
                            **document.metadata,
                        },
                    )
                )
                chunk_index += 1
                buffer = []
                buffer_len = 0

            buffer.append(para)
            buffer_len += len(para)

        if buffer:
            chunk_text = "\n".join(buffer).strip()
            chunks.append(
                TextChunk(
                    chunk_id=f"{document.doc_id}_chunk_{chunk_index:04d}",
                    doc_id=document.doc_id,
                    text=chunk_text,
                    chunk_index=chunk_index,
                    source_file=document.file_name,
                    section_title=None,
                    start_page=None,
                    end_page=None,
                    pages=[],
                    metadata={
                        "file_type": document.file_type,
                        **document.metadata,
                    },
                )
            )

        return chunks

    def _remove_repeated_page_lines(self, page_texts: List[dict]) -> List[dict]:
        if not page_texts:
            return []

        line_counter = Counter()
        total_pages = len(page_texts)
        threshold = max(3, total_pages // 3)

        for page in page_texts:
            lines = [x.strip() for x in page["text"].splitlines() if x.strip()]
            for line in set(lines):
                line_counter[line] += 1

        cleaned_pages = []
        for page in page_texts:
            lines = [x.strip() for x in page["text"].splitlines() if x.strip()]
            kept = []
            for line in lines:
                # 在很多页里都出现的短行，认为是页眉页脚，去掉
                if line_counter[line] >= threshold and len(line) <= 100:
                    continue
                kept.append(line)

            cleaned_pages.append(
                {
                    "page_num": page["page_num"],
                    "text": "\n".join(kept).strip(),
                }
            )

        return cleaned_pages

    def _page_to_paragraphs(self, page_text: str) -> List[str]:
        if not page_text:
            return []

        raw_paras = [p.strip() for p in re.split(r"\n\s*\n", page_text) if p.strip()]
        if raw_paras:
            return raw_paras

        lines = [x.strip() for x in page_text.splitlines() if x.strip()]
        if not lines:
            return []

        if not self.config.merge_short_paragraphs:
            return lines

        merged: List[str] = []
        buffer = ""
        for line in lines:
            if not buffer:
                buffer = line
            elif len(buffer) < self.config.min_chunk_chars:
                buffer += "\n" + line
            else:
                merged.append(buffer)
                buffer = line

        if buffer:
            merged.append(buffer)

        return merged

    def _make_doc_id(self, file_path: str | Path) -> str:
        p = Path(file_path)
        stem = re.sub(r"[^a-zA-Z0-9_]+", "_", p.stem).strip("_").lower()
        suffix = uuid.uuid4().hex[:8]
        return f"{stem}_{suffix}"
